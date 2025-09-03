import tkinter as tk
from tkinter import scrolledtext, PhotoImage, filedialog as fd, messagebox
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from PIL import Image, ImageTk

import threading
import queue
import requests
import time
import os
import re

# Optional dependency for Word doc generation
try:
    import docx
except ImportError:
    docx = None

# Optional deps for file parsing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pandas as pd
except ImportError:
    pd = None

# ---------------- Groq API Config ----------------
# API_KEY = os.getenv("GROQ_API_KEY", "your_default_key_here") # Replace with your key
API_KEY = os.getenv("GROQ_API_KEY")
if not API_KEY:
    raise ValueError("❌ GROQ_API_KEY is not set. Please set it as an environment variable.")

ENDPOINT = "https://api.groq.com/openai/v1/chat/completions"
MODEL = "llama-3.1-8b-instant"
MAX_CONTEXT_CHARS = 12000

reply_queue = queue.Queue()

# ---------------- Backend and Helper Functions (Unchanged) ----------------
def build_messages(user_message: str, file_context: str | None):
    system_prompt = (
        "You are Vishwasahayak, a futuristic AI assistant. "
        "Be concise, accurate, and reference the provided document context when relevant. "
        "The user may upload multiple files; their contents will be provided sequentially. "
        "When asked to compare, analyze, or summarize, clearly reference which file you are drawing information from.\n\n"
        "--- SPECIAL CAPABILITIES ---\n"
        "If the user asks you to 'extract', 'save', 'export', or 'download' your response, you can create a file.\n"
        "To do this, start your response with a special command on the very first line:\n"
        "1. For Word documents: `[SAVE_AS_DOCX: suggested_filename.docx]` followed by the text content.\n"
        "2. For Excel files: `[SAVE_AS_XLSX: suggested_filename.xlsx]` followed by the data formatted as a clean, pipe-delimited Markdown table.\n"
    )

    if file_context:
        context_preamble = (
            "The user has uploaded one or more files. The following is the extracted content from them, "
            "separated by headers. Use this content to answer the user's question. If information isn't present, say so.\n\n"
            "----- BEGIN EXTRACTED CONTEXT -----\n"
            f"{file_context}\n"
            "----- END EXTRACTED CONTEXT -----\n"
        )
        return [{"role": "system", "content": system_prompt}, {"role": "user", "content": context_preamble}, {"role": "user", "content": user_message}]
    else:
        return [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_message}]


def chat_with_groq(user_message, file_context=None):
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {"model": MODEL, "messages": build_messages(user_message, file_context), "temperature": 0.7}
    try:
        response = requests.post(ENDPOINT, headers=headers, json=payload, timeout=60)
        response_json = response.json()
        if response.status_code == 200 and "choices" in response_json:
            reply = response_json["choices"][0]["message"]["content"]
        else:
            error_message = response_json.get("error", {}).get("message", response.text)
            reply = f"⚠️ Error from API: {error_message}"
    except requests.exceptions.RequestException as e:
        reply = f"⚠️ Network Error: Could not connect to the API. {e}"
    except Exception as e:
        reply = f"⚠️ An unexpected error occurred: {e}"
    reply_queue.put(("Bot", reply))

def extract_word_text(path: str) -> str:
    if docx is None: raise RuntimeError("python-docx not installed.")
    document = docx.Document(path)
    return "\n".join([para.text for para in document.paragraphs]).strip()

def extract_pdf_text(path: str) -> str:
    if PyPDF2 is None: raise RuntimeError("PyPDF2 not installed.")
    text = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text.append(f"[Page {i+1}]\n{(page.extract_text() or '').strip()}\n")
    return "\n".join(text).strip()

def extract_excel_text(path: str) -> str:
    if pd is None: raise RuntimeError("pandas and openpyxl not installed.")
    sheets = pd.read_excel(path, sheet_name=None)
    blocks = [f"[Sheet: {name}]\n{df.fillna('').astype(str).iloc[:200, :30].to_string(index=False)}\n" for name, df in sheets.items()]
    return "\n".join(blocks).strip()

def save_as_word(content: str, path: str):
    if docx is None: raise RuntimeError("python-docx not installed.")
    doc = docx.Document(); doc.add_paragraph(content); doc.save(path)

def save_as_excel(markdown_table: str, path: str):
    if pd is None: raise RuntimeError("pandas and openpyxl not installed.")
    lines = [l for l in markdown_table.strip().split('\n') if not re.match(r'^\s*\|?s*:?-+:?s*\|?$', l.replace('|','').strip())]
    data = [[cell.strip() for cell in line.strip().strip('|').split('|')] for line in lines]
    if not data: raise ValueError("No data in Markdown table.")
    pd.DataFrame(data[1:], columns=data[0]).to_excel(path, index=False)

def smart_truncate(text: str, limit: int) -> str:
    if len(text) <= limit: return text
    cut = text[:limit]
    last_break = max(cut.rfind("\n\n"), cut.rfind("\n"))
    return cut[:last_break] + "\n\n[context truncated...]" if last_break > limit * 0.7 else cut + "\n\n[context truncated...]"

# ---------------- Modern Tkinter App Class ----------------
class ChatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🤖 Vishwasahayak")
        self.stop_typing = False
        self.is_typing = False
        self.active_context = ""
        self.file_names = []

        self._create_splash_screen()
        self._create_main_chat_ui()
        self.root.after(100, self.check_replies)

    def _create_splash_screen(self):
        self.root.geometry("700x500")
        self.splash_frame = ttk.Frame(self.root, padding=20)
        self.splash_frame.pack(fill=BOTH, expand=True)
        ttk.Label(self.splash_frame, text="🤖", font=("Segoe UI", 60), bootstyle="primary").pack(pady=10)
        ttk.Label(self.splash_frame, text="Welcome to Vishwasahayak", font=("Segoe UI", 28, "bold"), bootstyle="primary").pack(pady=5)
        self.description_text = ("🚀 The most futuristic AI assistant powered by Groq API,\n"
                                 "⚡ Experience lightning-fast responses that redefine AI speed and accuracy.\n"
                                 "🤝 Welcome to the future of AI interaction!")
        self.description_label = ttk.Label(self.splash_frame, text="", font=("Segoe UI", 12), bootstyle="secondary", justify="center")
        self.description_label.pack(pady=20)
        self.type_text(self.description_text, self.description_label)
        ttk.Button(self.splash_frame, text="🚀 Start Chat", bootstyle="success", command=self.show_chat_interface, padding=(20, 10)).pack(pady=20)

    def _create_main_chat_ui(self):
        self.main_chat_frame = ttk.Frame(self.root, padding=10)
        self.main_chat_frame.columnconfigure(0, weight=1); self.main_chat_frame.rowconfigure(1, weight=1)
        top_controls = ttk.Frame(self.main_chat_frame); top_controls.grid(row=0, column=0, sticky="ew", pady=(0, 10)); top_controls.columnconfigure(0, weight=1)
        self.ctx_status = ttk.Label(top_controls, text="No file context attached", font=("Segoe UI", 10), bootstyle="secondary"); self.ctx_status.grid(row=0, column=0, sticky="w")
        self.clear_ctx_btn = ttk.Button(top_controls, text="Clear Context", bootstyle="danger-outline", command=self.clear_context); self.clear_ctx_btn.grid(row=0, column=1, sticky="e")
        self.chat_display = scrolledtext.ScrolledText(self.main_chat_frame, wrap=tk.WORD, state="disabled", font=("Segoe UI", 12), relief=tk.FLAT, bd=0, padx=15, pady=15)
        self.chat_display.grid(row=1, column=0, sticky="nsew")
        input_frame = ttk.Frame(self.main_chat_frame, padding=(0, 10)); input_frame.grid(row=2, column=0, sticky="ew"); input_frame.columnconfigure(1, weight=1)
        try:
            upload_icon_img = Image.open("upload_icon.png").resize((20, 20), Image.LANCZOS)
            self.upload_icon = ImageTk.PhotoImage(upload_icon_img)
            self.upload_btn = ttk.Button(input_frame, image=self.upload_icon, bootstyle="secondary", command=self.upload_files); self.upload_btn.grid(row=0, column=0, sticky="w", padx=(0, 10))
        except FileNotFoundError:
            self.upload_btn = ttk.Button(input_frame, text="Attach", bootstyle="secondary", command=self.upload_files); self.upload_btn.grid(row=0, column=0, sticky="w", padx=(0, 10))
        self.entry = ttk.Entry(input_frame, font=("Segoe UI", 12), bootstyle="primary"); self.entry.grid(row=0, column=1, sticky="ew", ipady=5); self.entry.bind("<Return>", lambda e: self.send_or_stop())
        self.send_btn = ttk.Button(input_frame, text="Send", bootstyle="primary", command=self.send_or_stop); self.send_btn.grid(row=0, column=2, sticky="e", padx=(10, 0))
        self.chat_display.tag_config("user", foreground="#004085", font=("Segoe UI", 12, "bold"), spacing1=10, spacing3=10, lmargin1=10)
        self.chat_display.tag_config("bot", foreground="#333333", font=("Segoe UI", 12), spacing1=10, spacing3=10, lmargin1=10, background="#F1F1F1", lmargin2=20, rmargin=10, borderwidth=1, relief="solid")
        self.chat_display.tag_config("error", foreground="#721c24", font=("Segoe UI", 12, "italic"), background="#f8d7da")
        self.chat_display.tag_config("info", foreground="#0c5460", font=("Segoe UI", 10, "italic"), justify="center", spacing1=10)
        self.chat_display.tag_config("sender_user", font=("Segoe UI", 10, "bold"), foreground="#004085")
        self.chat_display.tag_config("sender_bot", font=("Segoe UI", 10, "bold"), foreground="#00796B")

    def type_text(self, text, label, index=0):
        if index <= len(text):
            label.config(text=text[:index])
            self.root.after(40, self.type_text, text, label, index + 1)

    def show_chat_interface(self):
        self.splash_frame.pack_forget()
        self.root.geometry("800x700")
        self.main_chat_frame.pack(fill=BOTH, expand=True)
        self.entry.focus_set()

    # --- CORRECTED: This function now appends files ---
    def upload_files(self):
        paths = fd.askopenfilenames(
            title="Select files to upload",
            filetypes=[
                ("All Supported Files", "*.pdf *.xlsx *.xls *.docx"),
                ("PDF Documents", "*.pdf"),
                ("Word Documents", "*.docx"),
                ("Excel Spreadsheets", "*.xlsx *.xls"),
            ]
        )
        if not paths:
            return

        self.add_message("Bot", f"📥 Processing {len(paths)} new file(s)...")
        
        newly_added_files = 0
        try:
            for path in paths:
                filename = os.path.basename(path)
                # Prevent adding the same file twice
                if filename in self.file_names:
                    self.add_message("Bot", f"ℹ️ Skipping duplicate file: {filename}")
                    continue

                ext = os.path.splitext(filename)[1].lower()
                
                file_header = f"\n\n--- START OF FILE: {filename} ---\n"
                
                if ext == ".pdf": extracted = extract_pdf_text(path)
                elif ext in (".xlsx", ".xls"): extracted = extract_excel_text(path)
                elif ext == ".docx": extracted = extract_word_text(path)
                else:
                    self.add_message("Bot", f"⚠️ Skipping unsupported file: {filename}")
                    continue

                if extracted.strip():
                    # Append new content and filename
                    self.active_context += (file_header + extracted)
                    self.file_names.append(filename)
                    newly_added_files += 1
                else:
                    self.add_message("Bot", f"⚠️ No text could be extracted from {filename}.")
            
            if newly_added_files == 0:
                self.add_message("Bot", "No new files were added.")
                return

            # Truncate the entire combined context
            full_context = self.active_context
            self.active_context = smart_truncate(full_context, MAX_CONTEXT_CHARS)

            badge = f"📎 Context: {len(self.file_names)} files attached"
            if len(self.file_names) <= 3:
                badge += f" ({', '.join(self.file_names)})"
            if len(full_context) > len(self.active_context):
                badge += " (truncated)"
            self.ctx_status.config(text=badge)

            self.add_message(
                "Bot",
                f"✅ {newly_added_files} new file(s) added successfully! "
                "The context now contains content from all uploaded documents."
            )

        except Exception as e:
            self.add_message("Bot", f"⚠️ An error occurred during file processing: {e}", typewriter=False)

    def clear_context(self):
        self.active_context = ""
        self.file_names = []
        self.ctx_status.config(text="No file context attached")
        self.add_message("Bot", "🧹 Cleared all attached file contexts.")

    def add_message(self, sender, message, typewriter=False):
        self.chat_display.config(state="normal")
        tag, sender_tag, sender_name = ("user", "sender_user", "🧑 You") if sender == "You" else ("bot", "sender_bot", "🤖 Vishwasahayak")
        if sender == "Bot":
            if "Error" in message or "⚠️" in message: tag = "error"
            elif "Processing" in message or "Cleared" in message or "processed" in message or "added" in message: tag = "info"
        self.chat_display.insert(tk.END, f"\n{sender_name}\n", (sender_tag,))
        if typewriter and sender == "Bot" and tag == "bot":
            self.is_typing = True; self.stop_typing = False; self.send_btn.config(text="Stop", bootstyle="danger")
            self.chat_display.see(tk.END); self.chat_display.update(); remaining_text = ""
            for i, char in enumerate(message):
                if self.stop_typing: remaining_text = message[i:]; break
                self.chat_display.insert(tk.END, char, tag); self.chat_display.see(tk.END); self.chat_display.update(); time.sleep(0.015)
            if remaining_text: self.chat_display.insert(tk.END, remaining_text, tag)
            self.chat_display.insert(tk.END, "\n"); self.is_typing = False; self.send_btn.config(text="Send", bootstyle="primary")
        else: self.chat_display.insert(tk.END, f"{message}\n", tag)
        self.chat_display.config(state="disabled"); self.chat_display.yview(tk.END)

    def send_or_stop(self):
        if self.is_typing: self.stop_typing = True
        else: self.send_message()

    def send_message(self):
        user_input = self.entry.get().strip()
        if not user_input: return
        self.entry.delete(0, tk.END)
        self.add_message("You", user_input)
        threading.Thread(target=chat_with_groq, args=(user_input, self.active_context or None), daemon=True).start()

    def handle_bot_response(self, message: str):
        match = re.match(r"\[SAVE_AS_(DOCX|XLSX):\s*(.+?)\s*\]\n?(.*)", message, re.DOTALL)
        if match:
            file_type, s_name, content = match.groups(); content = content.strip()
            ext, f_types, save_func = (".docx", [("Word", "*.docx")], save_as_word) if file_type=="DOCX" else (".xlsx", [("Excel", "*.xlsx")], save_as_excel)
            path = fd.asksaveasfilename(initialfile=s_name, defaultextension=ext, filetypes=f_types)
            if path:
                try: save_func(content, path); self.add_message("Bot", f"✅ File saved: {path}", False)
                except Exception as e: self.add_message("Bot", f"⚠️ Error saving file: {e}", False)
            if content: self.add_message("Bot", content, True)
        else: self.add_message("Bot", message, True)

    def check_replies(self):
        try:
            while True: self.handle_bot_response(reply_queue.get_nowait()[1])
        except queue.Empty: pass
        self.root.after(100, self.check_replies)

# ---------------- Main Execution ----------------
if __name__ == "__main__":
    root = ttk.Window(themename="flatly") 
    app = ChatApp(root)
    root.mainloop()